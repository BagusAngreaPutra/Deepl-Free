from __future__ import annotations

import os
import re
import hashlib
import json
import shutil
import socket
import sqlite3
import sys
import tempfile
import threading
import time
import zipfile
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass
from typing import Dict, List, Tuple

from lxml import etree

from translator_core_legacy import AM_TO_BR, apply_british_dictionary as legacy_apply_british_dictionary

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
TRANSLATABLE_PARTS = [
    'word/document.xml',
    'word/footnotes.xml',
    'word/endnotes.xml',
    'word/comments.xml',
]

# Google Translate accepts roughly 5,000 characters per request.  Keeping a
# little headroom lets us send substantially fewer requests without hitting
# that limit.  The old values made medium-sized documents take long enough for
# the web worker/proxy to close the connection (shown by fetch as
# "Failed to fetch").
BATCH_SIZE = 20
MAX_BATCH_CHARS = 4500
SEPARATOR = '\nZXQSEP001ZXQ\n'
PARALLEL_WORKERS = max(1, min(4, int(os.getenv('TRANSLATOR_PARALLEL_WORKERS', '1'))))
CACHE_VERSION = 'v2'

translate_cache: Dict[str, str] = {}
stats = {'translated': 0, 'cached': 0, 'skipped': 0, 'errors': 0}
_persistent_cache_db: sqlite3.Connection | None = None
_system_getaddrinfo = socket.getaddrinfo
_dns_cache: Dict[str, list] = {}
_http_sessions = threading.local()


def _http_session():
    """Reuse HTTPS connections per translation thread instead of exhausting Winsock."""
    import requests

    session = getattr(_http_sessions, 'session', None)
    if session is None:
        session = requests.Session()
        adapter = requests.adapters.HTTPAdapter(pool_connections=2, pool_maxsize=2, max_retries=0)
        session.mount('https://', adapter)
        _http_sessions.session = session
    return session


def _dns_cache_path() -> str:
    return os.getenv(
        'TRANSLATOR_DNS_CACHE_PATH',
        os.path.abspath('storage/app/private/translator-dns-cache.json'),
    )


def _serialise_addresses(addresses: list) -> list:
    return [
        [family, socktype, proto, canonname, list(sockaddr)]
        for family, socktype, proto, canonname, sockaddr in addresses
    ]


def _deserialise_addresses(addresses: list) -> list:
    return [
        (family, socktype, proto, canonname, tuple(sockaddr))
        for family, socktype, proto, canonname, sockaddr in addresses
    ]


def _save_dns_cache() -> None:
    cache_path = _dns_cache_path()
    os.makedirs(os.path.dirname(cache_path), exist_ok=True)
    temporary_path = cache_path + '.tmp'
    with open(temporary_path, 'w', encoding='utf-8') as handle:
        json.dump(
            {host: _serialise_addresses(value) for host, value in _dns_cache.items()},
            handle,
        )
    os.replace(temporary_path, cache_path)


def _resolve_ipv4_over_https(hostname: str) -> list:
    """Bootstrap DNS through a literal IP when the hosting resolver is down."""
    import requests

    errors = []
    for resolver_ip in ('1.1.1.1', '1.0.0.1'):
        try:
            response = requests.get(
                f'https://{resolver_ip}/dns-query',
                params={'name': hostname, 'type': 'A'},
                headers={
                    'Accept': 'application/dns-json',
                    'Host': 'cloudflare-dns.com',
                },
                timeout=(5, 10),
            )
            response.raise_for_status()
            payload = response.json()
            addresses = [
                answer['data'] for answer in payload.get('Answer', [])
                if answer.get('type') == 1 and answer.get('data')
            ]
            if addresses:
                emit_progress(
                    'translation_dns_bootstrapped',
                    hostname=hostname,
                    addresses=len(addresses),
                )
                return [
                    (socket.AF_INET, socket.SOCK_STREAM, socket.IPPROTO_TCP, '', (address, 443))
                    for address in addresses
                ]
            raise ValueError('DNS-over-HTTPS tidak mengembalikan alamat IPv4.')
        except Exception as resolver_error:
            errors.append(str(resolver_error)[:300])

    emit_progress('translation_dns_bootstrap_failed', hostname=hostname, errors=errors)
    return []


def install_resilient_dns(hostnames: Tuple[str, ...] = (
    'translate.google.com',
    'translate.googleapis.com',
    'clients5.google.com',
)) -> None:
    """Reuse last-known addresses when the Windows DNS resolver is intermittent."""
    global _dns_cache
    try:
        with open(_dns_cache_path(), encoding='utf-8') as handle:
            stored = json.load(handle)
        _dns_cache.update({host: _deserialise_addresses(value) for host, value in stored.items()})
    except (FileNotFoundError, OSError, ValueError, TypeError):
        pass

    cache_changed = False
    for hostname in hostnames:
        try:
            _dns_cache[hostname] = _system_getaddrinfo(hostname, 443, type=socket.SOCK_STREAM)
            cache_changed = True
        except socket.gaierror:
            if hostname not in _dns_cache:
                resolved = _resolve_ipv4_over_https(hostname)
                if resolved:
                    _dns_cache[hostname] = resolved
                    cache_changed = True
    if cache_changed:
        try:
            _save_dns_cache()
        except OSError as cache_error:
            emit_progress('translation_dns_cache_write_failed', error=str(cache_error)[:300])

    def resilient_getaddrinfo(host, port, *args, **kwargs):
        try:
            addresses = _system_getaddrinfo(host, port, *args, **kwargs)
            if host in hostnames:
                _dns_cache[host] = addresses
            return addresses
        except socket.gaierror:
            cached = _dns_cache.get(host)
            if not cached:
                raise
            emit_progress('translation_dns_cache_used', hostname=host)
            return [
                (family, socktype, proto, canonname, (sockaddr[0], port, *sockaddr[2:]))
                for family, socktype, proto, canonname, sockaddr in cached
            ]

    socket.getaddrinfo = resilient_getaddrinfo


def emit_progress(stage: str, **context) -> None:
    print(
        'JDS_PROGRESS '+json.dumps({'stage': stage, **context}, ensure_ascii=False),
        file=sys.stderr,
        flush=True,
    )


def persistent_cache_db() -> sqlite3.Connection:
    global _persistent_cache_db
    if _persistent_cache_db is None:
        cache_path = os.getenv(
            'TRANSLATOR_CACHE_PATH',
            os.path.abspath('storage/app/private/translator-cache.sqlite'),
        )
        os.makedirs(os.path.dirname(cache_path), exist_ok=True)
        _persistent_cache_db = sqlite3.connect(cache_path, timeout=15)
        _persistent_cache_db.execute('PRAGMA journal_mode=WAL')
        _persistent_cache_db.execute(
            'CREATE TABLE IF NOT EXISTS translations ('
            'cache_key TEXT PRIMARY KEY, translated_text TEXT NOT NULL, created_at TEXT DEFAULT CURRENT_TIMESTAMP)'
        )
    return _persistent_cache_db


def translation_cache_key(
    text: str,
    source_language: str,
    target_language: str,
    profile: str,
    custom_words: Dict[str, str] | None,
) -> str:
    custom_signature = repr(sorted((custom_words or {}).items()))
    raw = '\0'.join((CACHE_VERSION, source_language, target_language, profile, custom_signature, text))
    return hashlib.sha256(raw.encode('utf-8')).hexdigest()


def persistent_cache_get(cache_key: str) -> str | None:
    row = persistent_cache_db().execute(
        'SELECT translated_text FROM translations WHERE cache_key = ?', (cache_key,)
    ).fetchone()
    return row[0] if row else None


def persistent_cache_put(cache_key: str, translated_text: str) -> None:
    persistent_cache_db().execute(
        'INSERT OR REPLACE INTO translations (cache_key, translated_text) VALUES (?, ?)',
        (cache_key, translated_text),
    )


def persistent_cache_flush() -> None:
    if _persistent_cache_db is not None:
        _persistent_cache_db.commit()


@dataclass(frozen=True)
class Profile:
    key: str
    label: str
    description: str


PROFILES: Dict[str, Profile] = {
    'standard': Profile(
        key='standard',
        label='Standard British English',
        description='General British spelling and light polishing.',
    ),
    'academic': Profile(
        key='academic',
        label='British Academic English',
        description='Academic phrasing and cleaner university-style English.',
    ),
    'edu_academic': Profile(
        key='edu_academic',
        label='British Academic English for Course Documents',
        description='Optimised for CLO/Sub-CLO, lesson plans, tasks, and assessment rubrics.',
    ),
}


SOURCE_GLOSSARY_BASE: List[Tuple[str, str]] = [
    (r'\bkontrak kuliah\b', 'Course Agreement'),
    (r'\brencana pembelajaran semester\b', 'Semester Learning Plan'),
    (r'\bcapaian pembelajaran lulusan\b', 'Graduate Learning Outcomes'),
    (r'\bcapaian pembelajaran mata kuliah\b', 'Course Learning Outcomes'),
    (r'\bsub[- ]?capaian pembelajaran mata kuliah\b', 'Sub-course learning outcomes'),
    (r'\bcapaian pembelajaran\b', 'learning outcomes'),
    (r'\brubrik penilaian\b', 'Assessment Rubric'),
    (r'\bteknik penilaian\b', 'Assessment Method'),
    (r'\bpenugasan\b', 'assignment'),
    (r'\bbukti pengumpulan\b', 'evidence of submission'),
    (r'\bbukti penugasan\b', 'assessment evidence'),
    (r'\bpasangan minimal\b', 'minimal pairs'),
    (r'\bfitur artikulasi\b', 'articulatory features'),
    (r'\bfonetik\b', 'phonetics'),
    (r'\bfonologi\b', 'phonology'),
    (r'\bmorfologi\b', 'morphology'),
    (r'\bsintaksis\b', 'syntax'),
    (r'\bsemantik\b', 'semantics'),
    (r'\bpragmatik\b', 'pragmatics'),
    (r'\bsosiolinguistik\b', 'sociolinguistics'),
]

SOURCE_GLOSSARY_PROFILE: Dict[str, List[Tuple[str, str]]] = {
    'standard': SOURCE_GLOSSARY_BASE[:6],
    'academic': SOURCE_GLOSSARY_BASE[:12],
    'edu_academic': SOURCE_GLOSSARY_BASE,
}


START_VERB_FIXES = {
    'Shows': 'Explain',
    'Show': 'Explain',
    'Demonstrates': 'Demonstrate',
    'Demonstrating': 'Demonstrate',
    'Explores': 'Explore',
    'Exploring': 'Explore',
    'Compares': 'Compare',
    'Comparing': 'Compare',
    'Analyses': 'Analyse',
    'Analysing': 'Analyse',
    'Analyzes': 'Analyse',
    'Analyzing': 'Analyse',
    'Identifies': 'Identify',
    'Identifying': 'Identify',
    'Recognises': 'Recognise',
    'Recognising': 'Recognise',
    'Recognizes': 'Recognise',
    'Recognizing': 'Recognise',
    'Categorises': 'Categorise',
    'Categorising': 'Categorise',
    'Solving': 'Solve',
    'Discussing': 'Discuss',
    'Evaluating': 'Evaluate',
    'Evaluates': 'Evaluate',
    'Describes': 'Describe',
    'Describing': 'Describe',
    'Explains': 'Explain',
    'Explaining': 'Explain',
}


ACADEMIC_REPLACEMENTS: List[Tuple[str, str]] = [
    (r'\bsmallest pairs\b', 'minimal pairs'),
    (r'\barticulation features\b', 'articulatory features'),
    (r'\bcopy basic English sounds\b', 'transcribe basic English sounds'),
    (r'\busing IPA\b', 'using the IPA'),
    (r'\buse IPA\b', 'use the IPA'),
    (r'\bTuition Contract\b', 'Course Agreement'),
    (r'\bBill\s*\(proof\)', 'Evidence of submission'),
    (r'\bliveliness of answering directly\b', 'active participation and responsiveness during discussion'),
    (r'\bactive liveliness\b', 'active participation'),
    (r'\banswering directly\b', 'responding directly'),
    (r'\bthe smallest sound pair\b', 'the minimal pair'),
    (r'\bminimal pair[s]? smallest\b', 'minimal pairs'),
    (r'\blanguage problems in society\b', 'language-related problems in society'),
    (r'\bsociolinguistic variations\b', 'sociolinguistic variation'),
    (r'\bwhilst\b', 'while'),
    (r'\bAmongst\b', 'Among'),
    (r'\bamongst\b', 'among'),
    (r'\blicenced\b', 'licensed'),
    (r'\blicencing\b', 'licensing'),
    (r'\bprogramme\b(?=\s+code\b)', 'program'),
    (r'\bprogramme\b(?=\s+file\b)', 'program'),
]


EDU_SPECIFIC_REPLACEMENTS: List[Tuple[str, str]] = [
    (
        r'(?i)^\s*recogni[sz]e articulation features and copy basic English sounds using (?:the )?IPA\.?$',
        'Recognise articulatory features and transcribe basic English sounds using the IPA.',
    ),
    (
        r'(?i)^\s*show[s]? how meaning is derived from words and sentence structure\.?$',
        'Explain how meaning is derived from words and sentence structure.',
    ),
    (
        r'(?i)^\s*solve language(?:-related)? problems in society,? sociolinguistic variation[s]? and (?:its|their) implications\.?$',
        'Solve language-related problems in society and explain sociolinguistic variation and its implications.',
    ),
    (
        r'(?i)^\s*solve language(?:-related)? problems in society,? sociolinguistic variation and their implications\.?$',
        'Solve language-related problems in society and explain sociolinguistic variation and its implications.',
    ),
    (
        r'(?i)^\s*create a\s*(\d+)\s*[-–]?\s*(\d+)\s*(?:minutes?|minute)?\s+educational podcast\s+which\b',
        r'Create a \1–\2-minute educational podcast that',
    ),
    (
        r'(?i)^\s*create a\s*(\d+)\s*[-–]?\s*educational podcast\s*(\d+)\s*minutes?\s+that\b',
        r'Create a \1–\2-minute educational podcast that',
    ),
    (
        r'(?i)^\s*create a\s*(\d+)\s*[-–]?\s*educational podcast\s*(\d+)\s*minutes?\s+which\b',
        r'Create a \1–\2-minute educational podcast that',
    ),
    (
        r'(?i)^\s*create a\s*(\d+)\s*[-–]\s*educational podcast\s*(\d+)\s*minutes\s+which\b',
        r'Create a \1–\2-minute educational podcast that',
    ),
    (
        r'(?i)^\s*create a\s*(\d+)\s*to\s*(\d+)\s*minute educational podcast which\b',
        r'Create a \1–\2-minute educational podcast that',
    ),
    (
        r'(?i)^\s*able to analyse\b',
        'Analyse',
    ),
    (
        r'(?i)^\s*able to identify\b',
        'Identify',
    ),
    (
        r'(?i)^\s*able to explain\b',
        'Explain',
    ),
]


def parse_custom_words(raw_text: str) -> Dict[str, str]:
    custom: Dict[str, str] = {}
    for line in (raw_text or '').splitlines():
        line = line.strip()
        if not line or line.startswith('#'):
            continue
        if '=' in line:
            left, right = line.split('=', 1)
        elif ':' in line:
            left, right = line.split(':', 1)
        else:
            continue
        left = left.strip()
        right = right.strip()
        if not left or not right:
            continue
        custom[rf'\b{re.escape(left)}\b'] = right
    return custom


def apply_british_dictionary(text: str, custom_words: Dict[str, str] | None = None) -> str:
    if custom_words:
        original = dict(AM_TO_BR)
        try:
            AM_TO_BR.update(custom_words)
            text = legacy_apply_british_dictionary(text)
        finally:
            AM_TO_BR.clear()
            AM_TO_BR.update(original)
    else:
        text = legacy_apply_british_dictionary(text)

    for pattern, replacement in ACADEMIC_REPLACEMENTS:
        text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)

    # Normalise spacing/punctuation after replacements.
    text = re.sub(r'\s+([,.;:!?])', r'\1', text)
    text = re.sub(r'([({\[])\s+', r'\1', text)
    text = re.sub(r'\s+([)}\]])', r'\1', text)
    text = re.sub(r'\s{2,}', ' ', text).strip()
    return text


def is_translatable(text: str) -> bool:
    text = (text or '').strip()
    if not text or len(text) < 2:
        return False
    if re.fullmatch(r'[\d\s\.,;:!?@#$%^&*()\-_=+\[\]{}|<>/\\"\'"…•·°±×÷≤≥≠≈∞%]+', text):
        return False
    if re.match(r'https?://', text):
        return False
    if '@' in text and '.' in text and ' ' not in text:
        return False
    return bool(re.search(r'[A-Za-zÀ-ÿ]', text))


def protect_terms(text: str, profile: str) -> Tuple[str, Dict[str, str]]:
    glossary = SOURCE_GLOSSARY_PROFILE.get(profile, SOURCE_GLOSSARY_PROFILE['academic'])
    replacements: Dict[str, str] = {}
    protected_text = text
    index = 0

    for pattern, replacement in glossary:
        def _repl(match: re.Match[str]) -> str:
            nonlocal index
            token = f'ZXQTERM{index:03d}ZXQ'
            replacements[token] = replacement
            index += 1
            return token

        protected_text = re.sub(pattern, _repl, protected_text, flags=re.IGNORECASE)

    return protected_text, replacements


def restore_terms(text: str, replacements: Dict[str, str]) -> str:
    for token, replacement in replacements.items():
        text = text.replace(token, replacement)
    return text


def normalise_lo_lead(text: str) -> str:
    stripped = text.strip()
    if not stripped:
        return text

    # Short structured statements are likely CLO / rubric items.
    if len(stripped) > 220:
        return text

    stripped = re.sub(r'^(Students?|Learners?)\s+are\s+able\s+to\s+', '', stripped, flags=re.IGNORECASE)
    stripped = re.sub(r'^Able to\s+', '', stripped, flags=re.IGNORECASE)

    for src, tgt in START_VERB_FIXES.items():
        if stripped.startswith(src + ' '):
            return tgt + stripped[len(src):]

    return stripped


def polish_sentence_boundaries(text: str) -> str:
    text = re.sub(r'\s+([,.;:!?])', r'\1', text)
    text = re.sub(r'([,.;:!?])(\S)', r'\1 \2', text)
    text = re.sub(r'\s{2,}', ' ', text).strip()
    return text


def academic_polish(text: str, profile: str = 'academic') -> str:
    text = polish_sentence_boundaries(text)

    if profile in {'academic', 'edu_academic'}:
        text = re.sub(r'\bthe IPA\b', 'the IPA', text)
        text = re.sub(r'\b([0-9]+)\s*[-–]\s*([0-9]+)\s*(minutes?|minute|hours?|hour)\b', r'\1–\2-\3', text)
        text = re.sub(r'\b([0-9]+)\s+to\s+([0-9]+)\s*(minutes?|minute|hours?|hour)\b', r'\1–\2 \3', text)
        text = re.sub(r'\b([0-9]+)\s*[-–]\s*([0-9]+)-minutes\b', r'\1–\2-minute', text)
        text = re.sub(r'\b([0-9]+)\s*[-–]\s*([0-9]+)-minute\b', r'\1–\2-minute', text)
        text = re.sub(r'\b([0-9]+)\s*[-–]\s*([0-9]+)\s+minute\b', r'\1–\2-minute', text)
        text = re.sub(r'\b([0-9]+)\s*[-–]\s*([0-9]+)\s+minutes\b', r'\1–\2-minute', text)
        text = re.sub(r'\bwhich discusses and analyses\b', 'that discusses and analyses', text, flags=re.IGNORECASE)
        text = re.sub(r'\bwhich discusses\b', 'that discusses', text, flags=re.IGNORECASE)
        text = re.sub(r'\bwhich explains\b', 'that explains', text, flags=re.IGNORECASE)

    if profile == 'edu_academic':
        for pattern, replacement in EDU_SPECIFIC_REPLACEMENTS:
            text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
        text = normalise_lo_lead(text)
        for pattern, replacement in EDU_SPECIFIC_REPLACEMENTS:
            text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
        text = re.sub(
            r'(?i)\bminimal pairs? and articulatory features\b',
            'minimal pairs and articulatory features',
            text,
        )
        text = re.sub(
            r'(?i)\brespond directly in discussion\b',
            'respond appropriately during discussion',
            text,
        )

    if text and text[0].islower() and not text.startswith(('e.g.', 'i.e.')):
        text = text[0].upper() + text[1:]

    text = polish_sentence_boundaries(text)
    return text


def normalise_google_language(language: str) -> str:
    return 'en' if language in {'en-US', 'en-GB'} else language


def apply_output_style(
    text: str,
    target_language: str,
    profile: str,
    custom_words: Dict[str, str] | None = None,
) -> str:
    if target_language == 'en-GB':
        text = apply_british_dictionary(text, custom_words=custom_words)
    elif custom_words:
        for pattern, replacement in custom_words.items():
            text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
    return academic_polish(text, profile=profile) if target_language.startswith('en-') else text


def _get_translator(source_language: str = 'auto', target_language: str = 'en-GB'):
    from deep_translator import GoogleTranslator
    return GoogleTranslator(
        source=normalise_google_language(source_language),
        target=normalise_google_language(target_language),
    )


def _translate_with_google_fallbacks(
    text: str,
    source_language: str,
    target_language: str,
) -> str:
    """Try independent Google web endpoints when deep-translator is unavailable."""
    import requests

    language_params = {
        'sl': normalise_google_language(source_language),
        'tl': normalise_google_language(target_language),
        'q': text,
    }
    endpoints = (
        (
            'https://translate.googleapis.com/translate_a/single',
            {'client': 'gtx', 'dt': 't', **language_params},
            lambda payload: ''.join(
                segment[0] for segment in (payload[0] or [])
                if segment and segment[0]
            ),
        ),
        (
            'https://clients5.google.com/translate_a/t',
            {'client': 'dict-chrome-ex', **language_params},
            lambda payload: payload[0],
        ),
    )
    errors = []
    for url, params, extract_translation in endpoints:
        try:
            response = _http_session().get(url, params=params, timeout=(5, 30))
            response.raise_for_status()
            translated = extract_translation(response.json())
            if not isinstance(translated, str) or not translated:
                raise ValueError('Respons layanan terjemahan kosong atau tidak valid.')
            return translated
        except Exception as endpoint_error:
            errors.append(endpoint_error)

    # Preserve every endpoint error in diagnostics while keeping the last
    # network exception as the cause inspected by the retry classifier.
    emit_progress(
        'translation_fallbacks_failed',
        errors=[str(error)[:300] for error in errors],
    )
    raise errors[-1]


def translate_batch(batch_texts: List[str], source_language: str = 'auto', target_language: str = 'en-GB') -> List[str]:
    combined = SEPARATOR.join(batch_texts)
    try:
        translated = _translate_with_google_fallbacks(
            combined, source_language, target_language
        )
    except Exception as primary_error:
        if not is_translation_service_failure(primary_error):
            raise
        # Keep deep-translator as a final independent fallback, while the
        # pooled API connections remain the fast and stable normal path.
        translated = _get_translator(source_language, target_language).translate(combined)
    if SEPARATOR in translated:
        parts = translated.split(SEPARATOR)
    else:
        parts = translated.split('ZXQSEP001ZXQ')
    
    # Only strip newlines but preserve leading/trailing spaces from original context
    parts = [part.strip('\n\r\t') if part else '' for part in parts]
    while len(parts) < len(batch_texts):
        parts.append(batch_texts[len(parts)])
    return parts[:len(batch_texts)]


ENGLISH_MARKERS = {
    'the', 'and', 'of', 'to', 'in', 'is', 'are', 'for', 'with', 'that', 'this',
    'students', 'student', 'learning', 'course', 'assessment', 'outcomes', 'will',
}
INDONESIAN_MARKERS = {
    'dan', 'yang', 'di', 'ke', 'dari', 'untuk', 'dengan', 'adalah', 'ini', 'itu',
    'mahasiswa', 'pembelajaran', 'mata', 'kuliah', 'penilaian', 'capaian', 'akan',
}


def already_in_target_language(text: str, source_language: str, target_language: str) -> bool:
    if source_language != 'auto':
        return False
    words = re.findall(r"[a-zA-ZÀ-ÿ]+", text.lower())
    if len(words) < 4:
        return False
    english_score = sum(word in ENGLISH_MARKERS for word in words)
    indonesian_score = sum(word in INDONESIAN_MARKERS for word in words)
    if target_language.startswith('en-'):
        return english_score >= 2 and english_score >= indonesian_score * 2
    if target_language == 'id':
        return indonesian_score >= 2 and indonesian_score >= english_score * 2
    return False


def translate_batch_with_retry(
    batch_texts: List[str],
    source_language: str,
    target_language: str,
    attempts: int = 3,
) -> List[str]:
    last_error = None
    attempt = 0
    maximum_attempts = attempts
    while attempt < maximum_attempts:
        try:
            return translate_batch(batch_texts, source_language, target_language)
        except Exception as exc:
            last_error = exc
            # DNS on shared/local environments can disappear for several seconds.
            # Give network failures a longer recovery window without delaying
            # malformed-document or translation-content errors.
            if is_translation_service_failure(exc):
                maximum_attempts = max(maximum_attempts, 6)
            attempt += 1
            if attempt < maximum_attempts:
                delay = min(12.0, 0.75 * (2 ** (attempt - 1)))
                emit_progress(
                    'translation_batch_retry',
                    attempt=attempt,
                    maximum_attempts=maximum_attempts,
                    retry_in_seconds=delay,
                    error=str(exc)[:300],
                )
                time.sleep(delay)
    raise last_error


def is_translation_service_failure(error: Exception) -> bool:
    message = str(error).lower()
    markers = (
        'failed to resolve', 'nameresolutionerror', 'getaddrinfo failed',
        'connection refused', 'connection aborted', 'network is unreachable',
        'failed to establish a new connection', 'max retries exceeded',
        'requested service provider could not be loaded', 'winerror 10106',
        'too many requests', '429 client error',
    )
    return any(marker in message for marker in markers)


def build_translation_batches(
    items: List[Tuple[int, str, Dict[str, str], str]],
) -> List[List[Tuple[int, str, Dict[str, str], str]]]:
    batches = []
    cursor = 0
    while cursor < len(items):
        batch = []
        total_chars = 0
        while cursor < len(items) and len(batch) < BATCH_SIZE:
            item = items[cursor]
            projected = total_chars + len(item[1]) + len(SEPARATOR)
            if batch and projected > MAX_BATCH_CHARS:
                break
            batch.append(item)
            total_chars = projected
            cursor += 1
        if not batch:
            batch = [items[cursor]]
            cursor += 1
        batches.append(batch)
    return batches


def google_translate_paragraphs(
    texts: List[str],
    profile: str = 'academic',
    custom_words: Dict[str, str] | None = None,
    source_language: str = 'auto',
    target_language: str = 'en-GB',
) -> List[str]:
    install_resilient_dns()
    results = [''] * len(texts)
    to_translate: List[Tuple[int, str, Dict[str, str], str]] = []

    for i, text in enumerate(texts):
        if text in translate_cache:
            results[i] = translate_cache[text]
            stats['cached'] += 1
            continue
        if not is_translatable(text):
            results[i] = text
            stats['skipped'] += 1
            continue

        cache_key = translation_cache_key(text, source_language, target_language, profile, custom_words)
        cached = persistent_cache_get(cache_key)
        if cached is not None:
            translate_cache[text] = cached
            results[i] = cached
            stats['cached'] += 1
            continue

        if already_in_target_language(text, source_language, target_language):
            styled = apply_output_style(text, target_language, profile, custom_words)
            translate_cache[text] = styled
            results[i] = styled
            persistent_cache_put(cache_key, styled)
            stats['skipped'] += 1
            continue

        protected_text, replacements = protect_terms(text, profile) if target_language.startswith('en-') else (text, {})
        to_translate.append((i, protected_text, replacements, cache_key))

    batches = build_translation_batches(to_translate)
    emit_progress(
        'translation_batches_started',
        text_items=len(to_translate),
        batches=len(batches),
        parallel_workers=PARALLEL_WORKERS,
        cached=stats['cached'],
        skipped=stats['skipped'],
    )
    completed_batches = 0
    service_failure = None
    with ThreadPoolExecutor(max_workers=PARALLEL_WORKERS) as executor:
        future_batches = {
            executor.submit(
                translate_batch_with_retry,
                [item[1] for item in batch],
                source_language,
                target_language,
            ): batch
            for batch in batches
        }
        for future in as_completed(future_batches):
            batch = future_batches[future]
            try:
                parts = future.result()
                for (idx, _, replacements, cache_key), translated in zip(batch, parts):
                    restored = restore_terms(translated, replacements)
                    polished = apply_output_style(restored, target_language, profile, custom_words)
                    translate_cache[texts[idx]] = polished
                    results[idx] = polished
                    persistent_cache_put(cache_key, polished)
                    stats['translated'] += 1
            except Exception as batch_error:
                if is_translation_service_failure(batch_error):
                    service_failure = batch_error
                    for pending in future_batches:
                        pending.cancel()
                    break
                # Fallback per paragraph keeps one rejected batch from failing the document.
                for idx, original_src, replacements, cache_key in batch:
                    try:
                        translated = translate_batch_with_retry(
                            [original_src], source_language, target_language, attempts=2
                        )[0]
                        restored = restore_terms(translated, replacements)
                        polished = apply_output_style(restored, target_language, profile, custom_words)
                        translate_cache[texts[idx]] = polished
                        results[idx] = polished
                        persistent_cache_put(cache_key, polished)
                        stats['translated'] += 1
                    except Exception:
                        results[idx] = texts[idx]
                        stats['errors'] += 1

            completed_batches += 1
            report_every = max(1, len(batches) // 10)
            if completed_batches % report_every == 0 or completed_batches == len(batches):
                emit_progress(
                    'translation_batches_progress',
                    completed=completed_batches,
                    total=len(batches),
                    translated=stats['translated'],
                    errors=stats['errors'],
                )

    if service_failure is not None:
        emit_progress('translation_service_unavailable', error=str(service_failure)[:500])
        raise RuntimeError(
            'Layanan terjemahan tidak dapat dihubungi. Periksa koneksi internet atau DNS, '
            'lalu coba kembali beberapa saat lagi.'
        ) from service_failure

    persistent_cache_flush()

    return results


def get_paragraph_text(para_elem) -> str:
    return ''.join(
        (elem.text or '')
        for elem in para_elem.iter()
        if elem.tag == f'{{{W}}}t'
    )


def set_paragraph_text(para_elem, new_text: str):
    runs = [elem for elem in para_elem.iter() if elem.tag == f'{{{W}}}r']
    if not runs:
        return

    text_runs = []
    for run in runs:
        t_elems = run.findall(f'{{{W}}}t')
        if t_elems and any((t.text or '').strip() for t in t_elems):
            text_runs.append((run, t_elems))

    if not text_runs:
        return

    if len(text_runs) == 1:
        # Single run: put all text in first text element
        _, t_elems = text_runs[0]
        t_elems[0].text = new_text
        if new_text != new_text.strip():
            t_elems[0].set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        for t in t_elems[1:]:
            t.text = ''
        return

    # Multiple runs: put ALL translated text in the FIRST run to avoid
    # mid-word splits (e.g. "curr" + "iculum" or "a" + "nd").
    # Distributing by character proportion from the original language is
    # unreliable because translated text has different word boundaries.
    # Subsequent runs are cleared but kept in the XML so their formatting
    # marks (bold, italic, font size, etc.) remain untouched.
    _, first_t_elems = text_runs[0]
    first_t_elems[0].text = new_text
    if new_text != new_text.strip():
        first_t_elems[0].set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    for t in first_t_elems[1:]:
        t.text = ''

    # Clear text content from all subsequent runs (keep run element intact)
    for _, t_elems in text_runs[1:]:
        for t in t_elems:
            t.text = ''


def process_xml_part(
    xml_content: bytes,
    profile: str = 'academic',
    custom_words: Dict[str, str] | None = None,
    source_language: str = 'auto',
    target_language: str = 'en-GB',
) -> bytes:
    try:
        tree = etree.fromstring(xml_content)
    except etree.XMLSyntaxError:
        return xml_content

    paragraphs = tree.findall(f'.//{{{W}}}p')
    orig_texts = [get_paragraph_text(p) for p in paragraphs]
    translated = google_translate_paragraphs(
        orig_texts,
        profile=profile,
        custom_words=custom_words,
        source_language=source_language,
        target_language=target_language,
    )

    for para, new_text in zip(paragraphs, translated):
        if new_text and new_text != get_paragraph_text(para):
            set_paragraph_text(para, new_text)

    return etree.tostring(tree, xml_declaration=True, encoding='UTF-8', standalone=True)


def translate_pdf(
    input_path: str,
    output_path: str,
    output_format: str = 'pdf',
    custom_words: Dict[str, str] | None = None,
    profile: str = 'academic',
    source_language: str = 'auto',
    target_language: str = 'en-GB',
):
    """Translate text-layer PDF blocks into a layout-preserving PDF or editable DOCX."""
    global stats
    import fitz

    stats = {'translated': 0, 'cached': 0, 'skipped': 0, 'errors': 0}
    translate_cache.clear()
    profile = profile if profile in PROFILES else 'academic'
    output_format = output_format if output_format in {'pdf', 'docx'} else 'pdf'
    started_at = time.time()

    document = fitz.open(input_path)
    emit_progress('pdf_opened', pages=len(document), output_format=output_format)
    records = []
    page_sizes = []
    for page_index, page in enumerate(document):
        page_sizes.append((page.rect.width, page.rect.height))
        blocks = sorted(page.get_text('blocks'), key=lambda block: (round(block[1], 1), block[0]))
        for block in blocks:
            text = (block[4] or '').strip()
            block_type = block[6] if len(block) > 6 else 0
            if block_type != 0 or not is_translatable(text):
                continue
            records.append({
                'page': page_index,
                'rect': fitz.Rect(block[:4]),
                'text': text,
            })

    if not records:
        document.close()
        raise ValueError(
            'PDF tidak memiliki lapisan teks yang dapat diterjemahkan. '
            'PDF hasil scan perlu menjalani OCR terlebih dahulu.'
        )

    emit_progress('pdf_text_extracted', pages=len(page_sizes), text_blocks=len(records))

    translation_started = time.time()
    translated = google_translate_paragraphs(
        [record['text'] for record in records],
        profile=profile,
        custom_words=custom_words,
        source_language=source_language,
        target_language=target_language,
    )
    emit_progress(
        'pdf_translation_completed',
        elapsed_seconds=round(time.time() - translation_started, 2),
        translated=stats['translated'],
        cached=stats['cached'],
        skipped=stats['skipped'],
        errors=stats['errors'],
    )

    if output_format == 'pdf':
        emit_progress('pdf_output_started', pages=len(page_sizes))
        records_by_page = {}
        for record, translated_text in zip(records, translated):
            record['translated'] = translated_text
            records_by_page.setdefault(record['page'], []).append(record)

        for page_index, page_records in records_by_page.items():
            page = document[page_index]
            for record in page_records:
                page.add_redact_annot(record['rect'], fill=(1, 1, 1))
            page.apply_redactions()
            for record in page_records:
                rect = record['rect']
                text = record['translated']
                estimated_lines = max(1, text.count('\n') + 1)
                font_size = min(11.0, max(5.5, rect.height / (estimated_lines * 1.35)))
                while font_size >= 5.0:
                    spare = page.insert_textbox(
                        rect,
                        text,
                        fontsize=font_size,
                        fontname='helv',
                        color=(0, 0, 0),
                        lineheight=1.05,
                    )
                    if spare >= 0:
                        break
                    font_size -= 0.75
        document.save(output_path, garbage=4, deflate=True)
        document.close()
    else:
        from docx import Document
        from docx.shared import Inches, Pt

        emit_progress('docx_output_started', pages=len(page_sizes), text_blocks=len(records))
        output_document = Document()
        section = output_document.sections[0]
        if page_sizes:
            section.page_width = Inches(page_sizes[0][0] / 72)
            section.page_height = Inches(page_sizes[0][1] / 72)
        section.top_margin = section.bottom_margin = Inches(0.55)
        section.left_margin = section.right_margin = Inches(0.6)

        current_page = 0
        for record, translated_text in zip(records, translated):
            while record['page'] > current_page:
                output_document.add_page_break()
                current_page += 1
            paragraph = output_document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(5)
            run = paragraph.add_run(translated_text)
            run.font.name = 'Arial'
            run.font.size = Pt(10)
        output_document.save(output_path)
        document.close()

    emit_progress(
        'pdf_output_completed',
        output_format=output_format,
        output_size_bytes=os.path.getsize(output_path) if os.path.isfile(output_path) else None,
        elapsed_seconds=round(time.time() - started_at, 2),
    )

    return {
        'translated': stats['translated'],
        'cached': stats['cached'],
        'skipped': stats['skipped'],
        'errors': stats['errors'],
        'elapsed_seconds': time.time() - started_at,
        'pages': len(page_sizes),
        'text_blocks': len(records),
        'output_format': output_format,
        'source_language': source_language,
        'target_language': target_language,
        'parallel_workers': PARALLEL_WORKERS,
    }


def translate_docx_v3(
    input_path: str,
    output_path: str,
    custom_words: Dict[str, str] | None = None,
    profile: str = 'edu_academic',
    source_language: str = 'auto',
    target_language: str = 'en-GB',
):
    global stats
    stats = {'translated': 0, 'cached': 0, 'skipped': 0, 'errors': 0}
    translate_cache.clear()
    profile = profile if profile in PROFILES else 'edu_academic'

    start_total = time.time()
    tmp_dir = tempfile.mkdtemp()
    try:
        tmp_docx = os.path.join(tmp_dir, 'working.docx')
        shutil.copy2(input_path, tmp_docx)

        with zipfile.ZipFile(tmp_docx, 'r') as z:
            all_names = z.namelist()

        parts_to_process = list(TRANSLATABLE_PARTS)
        for name in all_names:
            if name.startswith('word/') and name.endswith('.xml'):
                base = os.path.basename(name)
                if (base.startswith('header') or base.startswith('footer')) and name not in parts_to_process:
                    parts_to_process.append(name)

        parts_found = [p for p in parts_to_process if p in all_names]
        translated_parts = {}
        emit_progress('docx_parts_started', total=len(parts_found))

        for part_index, part_name in enumerate(parts_found, start=1):
            emit_progress(
                'docx_part_started',
                current=part_index,
                total=len(parts_found),
                part=part_name,
            )
            with zipfile.ZipFile(tmp_docx, 'r') as z:
                xml_content = z.read(part_name)
            translated_parts[part_name] = process_xml_part(
                xml_content,
                profile=profile,
                custom_words=custom_words,
                source_language=source_language,
                target_language=target_language,
            )
            emit_progress('docx_part_completed', current=part_index, total=len(parts_found))

        tmp_output = os.path.join(tmp_dir, 'output.docx')
        emit_progress('docx_output_started')
        with zipfile.ZipFile(tmp_docx, 'r') as zin:
            with zipfile.ZipFile(tmp_output, 'w', compression=zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    if item.filename in translated_parts:
                        zout.writestr(item, translated_parts[item.filename])
                    else:
                        zout.writestr(item, zin.read(item.filename))

        shutil.copy2(tmp_output, output_path)
        emit_progress('docx_output_completed')
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)

    elapsed = time.time() - start_total
    return {
        'translated': stats['translated'],
        'cached': stats['cached'],
        'skipped': stats['skipped'],
        'errors': stats['errors'],
        'elapsed_seconds': elapsed,
        'output_path': output_path,
        'profile': profile,
        'source_language': source_language,
        'target_language': target_language,
        'parallel_workers': PARALLEL_WORKERS,
    }
